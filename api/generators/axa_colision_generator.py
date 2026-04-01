"""
Generador de informes AXA Colisión en formato .docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageOps
import io

class AXAColisionGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Configurar estilos del documento"""
        # Configurar márgenes y fuente por defecto
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _add_marco_teorico(self):
        """Agregar marco teórico (texto fijo de AXA)"""
        
        # MÉTODO DEDUCTIVO
        self.doc.add_heading('MÉTODO DEDUCTIVO:', level=2)
        self.doc.add_paragraph(
            'En este método se desciende de lo general a lo particular, de forma que '
            'partiendo de enunciados de carácter universal y utilizando instrumentos '
            'científicos, se infieren enunciados particulares, pudiendo ser axiomático-deductivo '
            'cuando las premisas de partida la constituyen axiomas (proposiciones no demostrables), '
            'o hipotético-deductivo si las premisas de partida son hipótesis contrastables.'
        )
        
        # MÉTODO INDUCTIVO
        self.doc.add_heading('MÉTODO INDUCTIVO:', level=2)
        self.doc.add_paragraph(
            'El método inductivo es aquel método científico que alcanza conclusiones generales '
            'partiendo de hipótesis o antecedentes en particular, por lo tanto se puede decir '
            'que asciende de lo particular a lo general.'
        )
        
        # OBSERVACIÓN DIRECTA
        self.doc.add_heading('OBSERVACIÓN DIRECTA:', level=2)
        self.doc.add_paragraph(
            'Es una técnica que consiste en observar atentamente el fenómeno, hecho o caso, '
            'tomar información y registrarla para su posterior análisis.'
        )
        
        # MÉTODO CIENTÍFICO
        self.doc.add_heading('MÉTODO CIENTÍFICO:', level=2)
        self.doc.add_paragraph(
            'El método científico guía y ayuda a comprender cosas desconocidas por medio de '
            'la aplicación sistemática de sus pasos fundamentales:'
        )
        
        # PRINCIPIOS
        self.doc.add_heading('PRINCIPIOS', level=2)
        
        self.doc.add_paragraph('Principio de intercambio:', style='List Number')
        self.doc.add_paragraph(
            'En la comisión de un delito, el autor deja indicios de su parte y a la vez, '
            'arrastra con otros de los hechos.'
        )
        
        self.doc.add_paragraph('Principio de correspondencia:', style='List Number')
        self.doc.add_paragraph('Establece la relación de indicios con el autor del hecho.')
        
        self.doc.add_paragraph('Principio de reconstrucción:', style='List Number')
        self.doc.add_paragraph(
            'Permite deducir de los indicios recogidos en el lugar de la investigación, '
            'la forma en que ocurrió el hecho.'
        )
        
        self.doc.add_paragraph('Principio de Probabilidad:', style='List Number')
        self.doc.add_paragraph(
            'Permite deducir la probabilidad o imposibilidad de un fenómeno con base en '
            'el número de características verificadas durante un cotejo.'
        )
        
        # CONDICIONES GENERALES DE LA PÓLIZA
        self.doc.add_heading('CONDICIONES GENERALES DE LA PÓLIZA', level=2)
        self.doc.add_paragraph('Cláusula 12a. Pérdida del Derecho a Ser Indemnizado')
        self.doc.add_paragraph('Las obligaciones de la Compañía quedarán extinguidas:')
        
        self.doc.add_paragraph(
            '1. Si se demuestra que el Asegurado, el Conductor, el Beneficiario o sus '
            'Representantes, con el fin de hacer incurrir a la Compañía en error, disimulan '
            'o declaran inexactamente hechos que excluirían o podrían restringir dichas obligaciones.'
        )
        
        self.doc.add_paragraph(
            '2. Si en el Siniestro hubiere dolo o mala fe del Asegurado, del Conductor, '
            'del Beneficiario o de sus respectivos causahabientes.'
        )
        
        # LEY DEL CONTRATO DEL SEGURO
        self.doc.add_heading('LEY DEL CONTRATO DEL SEGURO', level=2)
        
        self.doc.add_heading('ARTICULO 69.', level=3)
        self.doc.add_paragraph(
            'La empresa aseguradora tendrá el derecho de exigir del asegurado o beneficiario '
            'toda clase de informaciones sobre los hechos relacionados con el siniestro y '
            'por los cuales puedan determinarse las circunstancias de su realización y las '
            'consecuencias del mismo.'
        )
        
        self.doc.add_heading('ARTICULO 70.', level=3)
        self.doc.add_paragraph(
            'Las obligaciones de la empresa quedarán extinguidas si demuestra que el asegurado, '
            'el beneficiario o los representantes de ambos, con el fin de hacerla incurrir en error, '
            'disimulan o declaran inexactamente hechos que excluirían o podrían restringir '
            'dichas obligaciones.'
        )
        
        # CÓDIGO PENAL
        self.doc.add_heading('CÓDIGO PENAL', level=2)
        self.doc.add_paragraph(
            'Artículo 386.- Comete el delito de fraude el que engañando a uno o aprovechándose '
            'del error en que éste se halla se hace ilícitamente de alguna cosa o alcanza un '
            'lucro indebido.'
        )
        
        self.doc.add_page_break()
    
    def _add_tabla_siniestro(self, datos_caso, analisis_ia):
        """Agregar tabla con datos del siniestro"""
        
        # Crear tabla 2x2
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Fila 1: Tipo de Siniestro
        cells = table.rows[0].cells
        cells[0].text = 'Tipo de Siniestro:'
        cells[1].text = analisis_ia.get('descripcion_hechos', '')
        
        # Fila 2: Fecha, Hora y Lugar (header)
        cells = table.rows[1].cells
        cells[0].text = 'Fecha, Hora y Lugar:'
        cells[1].text = 'Fecha, Hora y Lugar:'
        
        # Fila 3: Fecha y hora
        cells = table.rows[2].cells
        cells[0].text = 'Fecha y hora:'
        cells[1].text = f"{datos_caso.get('fecha_siniestro', '')} a las {datos_caso.get('hora_siniestro', '')}"
        
        # Fila 4: Lugar
        cells = table.rows[3].cells
        cells[0].text = 'Lugar:'
        cells[1].text = analisis_ia.get('lugar_hechos', '')
        
        self.doc.add_paragraph()
    
    def _add_tabla_conductor(self, datos_caso):
        """Agregar tabla con datos del conductor"""
        
        self.doc.add_heading('DATOS DEL CONDUCTOR ASEGURADO', level=3)
        
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = 'DATOS DEL CONDUCTOR ASEGURADO.'
        cells[1].text = 'DATOS DEL CONDUCTOR ASEGURADO.'
        
        cells = table.rows[1].cells
        cells[0].text = 'Nombre completo:'
        cells[1].text = datos_caso.get('nombre_conductor', '')
        
        self.doc.add_paragraph()
    
    def _add_tabla_vehiculo(self, datos_caso):
        """Agregar tabla con datos del vehículo"""
        
        self.doc.add_heading('DATOS DEL VEHÍCULO ASEGURADO', level=3)
        
        table = self.doc.add_table(rows=9, cols=2)
        table.style = 'Table Grid'
        
        # Header
        cells = table.rows[0].cells
        cells[0].text = 'DATOS DEL VEHÍCULO ASEGURADO.'
        cells[1].text = 'DATOS DEL VEHÍCULO ASEGURADO.'
        
        # Datos
        datos_vehiculo = [
            ('Tipo de bien:', 'Vehículo'),
            ('Marca:', datos_caso.get('marca', '')),
            ('Submarca:', datos_caso.get('submarca', '')),
            ('Modelo:', datos_caso.get('modelo', '')),
            ('Color:', datos_caso.get('color', '')),
            ('Placas:', datos_caso.get('placas', '')),
            ('Número de serie:', datos_caso.get('no_serie', '')),
            ('Uso:', datos_caso.get('uso', 'Particular'))
        ]
        
        for i, (label, value) in enumerate(datos_vehiculo, start=1):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = str(value)
        
        self.doc.add_paragraph()
    
    def _add_lista_danos(self, lista_danos):
        """Agregar tabla con lista de daños"""
        
        self.doc.add_heading('IDENTIFICACIÓN DE DAÑOS', level=3)
        
        # Crear tabla de 1 celda con todo el texto
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        cell = table.rows[0].cells[0]
        
        # Texto introductorio
        intro = (
            'La unidad marca {marca}, tipo {submarca}, modelo {modelo}, con placas de '
            'circulación {placas} y número de serie {no_serie} fue revisada de modo sensorial, '
            'utilizando la técnica de cronos y el procedimiento deductivo de lo general a lo '
            'particular, iniciando por la parte frontal y dividiendo el vehículo en líneas '
            'imaginarias (división segmentaria) y siguiendo el sentido horario, presentando '
            'los siguientes daños:\n\n'
        )
        
        # Agregar lista de daños
        texto_completo = intro + '\n'.join(lista_danos)
        cell.text = texto_completo
        
        self.doc.add_paragraph()
    
    def _add_image_with_border(self, image_path_or_bytes, width_inches=5):
        """Agregar imagen con marco negro de 2px"""
        
        # Abrir imagen
        if isinstance(image_path_or_bytes, bytes):
            img = Image.open(io.BytesIO(image_path_or_bytes))
        else:
            img = Image.open(image_path_or_bytes)
        
        # Agregar marco negro de 2px
        img_with_border = ImageOps.expand(img, border=2, fill='black')
        
        # Guardar en buffer
        img_buffer = io.BytesIO()
        img_with_border.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        # Agregar al documento
        self.doc.add_picture(img_buffer, width=Inches(width_inches))
        
        # Centrar
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate(self, datos_caso, analisis_ia, imagenes=None):
        """
        Generar informe completo de AXA Colisión
        
        Args:
            datos_caso: Dict con datos del formulario
            analisis_ia: Dict con análisis generado por Claude
            imagenes: List de bytes de imágenes a insertar
        
        Returns:
            BytesIO con el documento Word generado
        """
        
        # 1. Marco teórico
        self._add_marco_teorico()
        
        # 2. Datos del siniestro
        self._add_tabla_siniestro(datos_caso, analisis_ia)
        
        # 3. Descripción del lugar
        if 'descripcion_lugar' in analisis_ia:
            self.doc.add_paragraph(analisis_ia['descripcion_lugar'])
            self.doc.add_paragraph()
        
        # 4. Mapa/foto del lugar (si hay)
        if imagenes and len(imagenes) > 0:
            self.doc.add_paragraph('Ilustración 1.- Plano de ubicación del lugar de los hechos')
            self._add_image_with_border(imagenes[0])
            self.doc.add_paragraph()
        
        # 5. Datos del conductor
        self._add_tabla_conductor(datos_caso)
        
        # 6. Datos del vehículo
        self._add_tabla_vehiculo(datos_caso)
        
        # 7. Lista de daños
        lista_danos = analisis_ia.get('lista_danos', [])
        if lista_danos:
            self._add_lista_danos(lista_danos)
        
        # 8. Fotos de daños
        if imagenes and len(imagenes) > 1:
            self.doc.add_heading('FIJACIONES FOTOGRÁFICAS', level=3)
            for i, img_bytes in enumerate(imagenes[1:], start=2):
                self.doc.add_paragraph(f'Fotografía {i}')
                self._add_image_with_border(img_bytes, width_inches=4)
                self.doc.add_paragraph()
        
        # 9. Análisis y conclusiones
        if 'analisis' in analisis_ia:
            self.doc.add_page_break()
            self.doc.add_heading('ANÁLISIS TÉCNICO', level=2)
            self.doc.add_paragraph(analisis_ia['analisis'])
        
        if 'conclusiones' in analisis_ia:
            self.doc.add_heading('CONCLUSIONES', level=2)
            self.doc.add_paragraph(analisis_ia['conclusiones'])
        
        # Guardar en buffer
        doc_buffer = io.BytesIO()
        self.doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer

# Función auxiliar para usar desde Node.js
def generate_axa_colision_docx(datos_caso, analisis_ia, imagenes=None):
    """Wrapper para llamar desde Node.js"""
    generator = AXAColisionGenerator()
    return generator.generate(datos_caso, analisis_ia, imagenes)

